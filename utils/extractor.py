import io
import PyPDF2
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes.
    """
    pdf_file = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes.
    """
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    text = ""
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text += paragraph.text + "\n"
    # Extract text from tables if present
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text]
            text += " | ".join(row_text) + "\n"
    return text.strip()

def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extracts text from raw bytes with fallback encodings.
    """
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file using utf-8, latin-1, or cp1252.")
